# ecoding:utf-8
# 不加上面这个突然就有问题了，郁闷,由于内网的matplotlib包不能使用，安装又非常麻烦，所以我们更换包到pyecharts
from docx import Document  # 需要初始化python_docx
from docx.oxml.ns import qn
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT  # 设置表格对齐方式
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches
from docx.shared import RGBColor
from sqlcommand.yb_sql import sqlstr
from sqlcommand.getsqldata import getsqldata
from functools import reduce
from matplotlib import pyplot as plt
from pyecharts.charts import Map
from collections import defaultdict
from pyecharts import options as opts
from pyecharts.render import make_snapshot
from pyecharts.charts import Line
from pyecharts.charts import Pie
from pyecharts.charts import Bar
import calendar
import time
from pyecharts.globals import CurrentConfig  # 用于设置pyecharts读取本地js，而不会再去网上下载
import xlsxwriter
from wbnr.yb_text import get_ybwb
from snapshot_selenium import snapshot
import os
import shutil

'''
# 1.xlwt 只能创建拓展名为 xls 的 excel 文件, 无法创建 xlsx 拓展名, 而 xlsxwriter 两个都可以创建
# 2.xlwt 所用样式过多会导致无法使用,xlsxwriter 则不会有问题
# 3. 当处理巨大表格的时候建议使用 xlsxwriter,xlwt 生成的 xls 文件最多只能支持 65536 行数据, 而 xlsx 文件最大能够支持 1048576 行,16384 列数据
为了简化不同数据类型的写入过程，write方法已作为其他更加具体的数据类型方法的别名，如：

  write.string()：写入字符串类型数据
  wirte_number()：写入数字型数据
  write_blank()：写入空类型数据
  write_formula()：写入公式型数据
  write_datetime()：写入日期型数据
  wirte_boolean()：写入逻辑型数据
  write_url()：写入超链接型数据
'''
'''
1.目录  # 暂时还没有头绪难点也是无法获取页码的问题
2 将生成好的Excel对象插入word 没有头绪
3 页脚插入页码
'''


class scword():
    """
    功能：用于生成组内月报，这个对于周期性的文档报告很有意义
    author:aaa
    date:20200527
    标准：大标题23 一级标题22 二级标题15 三级标题12 一般字体直接输出，行距1.5倍，段前0.5行 段后0.5行 中文：宋体 英文：Times New Roman
    实现思路：三级架构，准备好每个一级目录下的二级目录和二级所对应的内容对象，然后统一按照结构循环输出
    问题：1 页眉图片未插入，2 Excel如何插入进来 3 目录如何自动生成 4 Excel如何自动使用列宽 5 页脚的页码如何插入
    """

    def __init__(self, ny):
        self.ny = ny
        self.wdmc = r'..\appendix\GT3-SJZY-YPT-云平台数据管理子项目-数据集成运维报告-{ny}.docx'.format(ny=ny)  # 要求必须传入文档名称

    doc = Document()  # 生成一个word对象
    # 设置pyecharts读取的路径
    js_path = r"../js/"
    CurrentConfig.ONLINE_HOST = js_path

    def checkdir(self):
        '''用于将历史生成的文件放到 old 文件夹，让附件目录一进去只看到当前生成的 '''
        curdir = os.path.abspath("../appendix")
        olddir = os.path.join(curdir, 'old')
        print("原始目录是{curdir}目标目录是：{olddir}".format(olddir=olddir, curdir=curdir))
        if not os.path.exists(olddir):
            os.mkdir(olddir)
            print("{olddir}文件不存在，已创建".format(olddir=olddir))
        # 获取之前生成的文件移动到old中
        files = os.listdir(curdir)
        print(files)
        for file in files:
            filepath = os.path.join(curdir, file)
            oldpath = os.path.join(olddir, file)
            if os.path.isfile(filepath):
                dqsj = time.strftime('%Y%m%d%H%M%S', time.localtime(time.time()))
                print("正在移动：", filepath)
                if os.path.exists(oldpath):
                    os.rename(oldpath, oldpath + dqsj)
                shutil.move(filepath, olddir)

    def setfooter(self):
        """处理页眉页脚 目前没有用，虽然新增了制定的页数什么都对，但是问题就是内容没有写在这些页面，也不知道那些内容在那些页"""
        for i in range(24):
            self.doc.add_section()  # 每次进来先新增一个
            # 添加页眉页脚
            header = self.doc.sections[i].header  # 获取当前的节
            header.is_linked_to_previous = False  # 断开上一节为了设置不同的页脚
            ymdl = header.paragraphs[0]  # 获取段落
            # 页眉段落第一个块添加图片靠左对齐，第二个块靠右对齐
            run1 = ymdl.add_run()
            run1.add_picture(r"../img/img_ym01.png", width=Cm(2.8), height=Cm(0.9))
            run1.add_picture(r"../img/img_ym02.png", width=Cm(1.3), height=Cm(0.9))
            run1.underline = True
            # 添加第二个块
            run2 = ymdl.add_run()
            run2.add_text("\t金税三期工程第二阶段总局数据资源建设项目云平台数据管理子项目")
            font = run2.font
            font.size = Pt(10.5)
            run2.underline = True
            # 段落样式的设定
            ymdl.style = self.doc.styles["Header"]
            pf = ymdl.paragraph_format
            pf.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            footer = self.doc.sections[i].footer
            footer.is_linked_to_previous = False
            paragraph = footer.paragraphs[0]
            zys = '24'
            yj_text = "第 {dqy} 页 共 {zys} 页".format(dqy=i + 1, zys=zys)
            paragraph.add_run(yj_text)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            i += 1

    def scword(self):
        '''生成文档'''
        self.checkdir()  # 处理目录
        self.setfooter()  # 设置页眉页脚
        print("历史附件已经移动到old文件夹中")
        # 文档对象因为内容中也要获取，所以增加到公共部分
        # 公共部分
        doc = self.doc
        nf = int(self.ny[:4])
        yf = int(self.ny[-2:])
        start, end = calendar.monthrange(nf, yf)
        table_style = "Table Grid"
        # # 设置一个空白样式 整体样式
        style = doc.styles['Normal']
        # # 设置西文字体
        style.font.name = u'Times New Roman'
        # # 设置中文字体
        style.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        ################################## 标题和前置表格 ##################################

        # 大标题
        paragraph = doc.add_heading(level=0)  # 初始化一个段落对象
        dbt = paragraph.add_run('''
        \n金税三期第二阶段总局数据资源建设项目\n云平台数据管理子项目\n数据集中链路报告\nV0.1.00\n\n\n
        ''')
        # 添加run对象，参数为text=None和style=None,
        # run对象有bold（加粗）和italic（斜体）这两个属性
        dbt.bold = True
        dbt.font.name = u'宋体'
        dbt._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        # 字体对象
        btzt = dbt.font
        btzt.size = Pt(23)  # 设置标题字体大小
        # btzt.bold=True
        btzt.color.rgb = RGBColor(0, 0, 0)
        # btzt.size=Inches(1)  也是设置字体大小，英寸
        # 设置标题水平居中
        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # 添加一个文档说明表格 table_style
        wdsmbg = doc.add_table(rows=7, cols=4, style=table_style)
        wdsmbg.cell(0, 0).text = '文档敏感性定义'
        wdsmbg.cell(0, 1).text = '敏感'
        wdsmbg.cell(0, 1).merge(wdsmbg.cell(0, 2)).merge(wdsmbg.cell(0, 3))  # 合并后面两个格子的对象

        wdsmbg.cell(1, 0).text = '编写人'
        wdsmbg.cell(1, 1).text = r'个人/组'
        wdsmbg.cell(1, 2).text = '编写日期'
        wdsmbg.cell(1, 3).text = r'YYYY-MM-DD'

        wdsmbg.cell(2, 0).text = '审核人'
        wdsmbg.cell(2, 1).text = '个人'
        wdsmbg.cell(2, 2).text = '审核日期'
        wdsmbg.cell(2, 3).text = r'YYYY-MM-DD'

        wdsmbg.cell(3, 0).text = '公开范围'
        wdsmbg.cell(3, 1).text = '国家税务总局、云平台管理项目组'
        wdsmbg.cell(3, 1).merge(wdsmbg.cell(3, 2)).merge(wdsmbg.cell(3, 3))

        wdsmbg.cell(4, 0).text = '建设单位'
        wdsmbg.cell(4, 1).text = '国家税务总局'
        wdsmbg.cell(4, 1).merge(wdsmbg.cell(4, 2)).merge(wdsmbg.cell(4, 3))

        wdsmbg.cell(5, 0).text = '承建单位'
        wdsmbg.cell(5, 1).text = '中国软件与技术服务股份有限公司'
        wdsmbg.cell(5, 1).merge(wdsmbg.cell(5, 2)).merge(wdsmbg.cell(5, 3))

        wdsmbg.cell(6, 0).text = '监理单位'
        wdsmbg.cell(6, 1).text = '北京赛迪工业和信息化工程监理中心有限公司'
        wdsmbg.cell(6, 1).merge(wdsmbg.cell(6, 2)).merge(wdsmbg.cell(6, 3))

        # 将表格内容设置为居中
        for r in range(7):
            for c in range(4):
                wdsmbg.cell(r, c).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        # 修订表格标题
        paragraph = doc.add_paragraph()  # 初始化一个段落对象
        xdbt = paragraph.add_run("\n\n\n修订状况\n")
        xdbtzt = xdbt.font
        xdbtzt.size = Pt(18)  # 设置标题字体大小
        xdbtzt.bold = True
        # 设置标题水平居中
        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # 修订内容表格
        xdbg = doc.add_table(rows=5, cols=6, style=table_style)
        xdbg.cell(0, 0).text = "章节编号"
        xdbg.cell(0, 1).text = "章节名称"
        xdbg.cell(0, 2).text = "修订内容简述"
        xdbg.cell(0, 3).text = "修订人"
        xdbg.cell(0, 4).text = "修订日期"
        xdbg.cell(0, 5).text = "批准人"
        # 设置标题居中
        for c in range(6):
            cp = xdbg.cell(0, c).paragraphs[0]
            cp.paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        # doc.add_page_break()  # 新开一个页面

        ################################## 输出文档标题和前置表格结束 ##################################

        # 目录部分
        paragraph = doc.add_paragraph()  # 初始化一个段落对象
        xdbt = paragraph.add_run("目 录\n")
        xdbtzt = xdbt.font
        xdbtzt.size = Pt(18)  # 设置标题字体大小
        xdbtzt.bold = True
        # 设置标题水平居中
        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # 设置段落格式
        paragraph_format.space_before = Pt(5)  # 上行间距
        paragraph_format.space_after = Pt(10)  # 下行间距
        paragraph_format.line_spacing = Pt(20)  # 行距

        # 获取到固定内容和动态内容的查询sql对象
        ml, nr2, nr3, zj_2 = get_ybwb.hqnr(get_ybwb())
        sql = sqlstr(self.ny)
        print("正在生成的是{ny}的数据，请注意核实".format(ny=sql.ny))

        # 一级目录加粗，二级目录直接输出,第二章的三级目录单独处理下,只要目录页面
        for dl in ml.keys():
            mldl = (doc.add_paragraph()).add_run(dl)
            (mldl.font).bold = True
            if dl[:dl.index("\t")] == "第6章":
                break
            elif dl[:dl.index("\t")] == "第2章":
                for ej in ml.get(dl):
                    # for ejbt in ej.keys():
                    #     # print("第二级标题",ejbt)
                    p2 = doc.add_paragraph(ej.keys())  # 直接输出
                    pf2 = p2.paragraph_format
                    pf2.first_line_indent = Inches(0.3)  # 首行缩进
                    for ejz in ej.values():
                        for sj in ejz:
                            p23 = doc.add_paragraph(sj.keys())  # 直接输出
                            pf23 = p23.paragraph_format
                            pf23.first_line_indent = Inches(0.6)  # 缩进双倍

            else:
                for xl in ml.get(dl):
                    for ej in xl.keys():
                        paragraph = doc.add_paragraph(ej)  # 直接输出
                        paragraph_format = paragraph.paragraph_format
                        paragraph_format.first_line_indent = Inches(0.3)  # 首行缩进
        doc.add_page_break()
        ################################## 目录结束 ##################################

        # 添加正文内容所有目录，目录和内容全部都要输出
        for dl in ml.keys():
            dldl = doc.add_heading(level=1)  # 添加标题
            pf = dldl.paragraph_format
            pf.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = dldl.add_run(dl[:dl.rindex("\t")])
            run.font.name = u'宋体'  # 由于heading修改了字体样式，但是中文标题这个样式会有问题，所以进行调整
            run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            dlbtzt = run.font
            # dlbtzt.bold = True
            dlbtzt.size = Pt(22)

            # 设置颜色，这两种都可以
            dlbtzt.color.rgb = RGBColor(0, 0, 0)  # WD_COLOR_INDEX.AUTO
            # dlbtzt.color.theme_color=WD_COLOR_INDEX.BLACK

            # 一些不同格式的处理

            if dl[:dl.index("\t")] == "第3章":
                dl3 = doc.add_paragraph(nr3)
                pf3 = dl3.paragraph_format
                pf3.first_line_indent = Inches(0.3)
                dl3.add_run().add_picture(r"..\img\img_dl3.png", width=Inches(6.25))
                # 设置段落格式
                pf3.space_before = Pt(5)  # 上行间距
                pf3.space_after = Pt(10)  # 下行间距
                pf3.line_spacing = Pt(20)  # 行距
                # pf3.left_indent = Pt(10)  左右距离不需要默认是0
                # pf3.right_indent = Pt(5)
            elif dl[:dl.index("\t")] == "第6章":
                break

            if dl[:dl.index("\t")] == "第2章":
                # 处理大标题直属内容
                doc.add_paragraph("\n")
                doc.add_picture(r"..\img\img_dl2.jpg", width=Inches(6.25))  # 这个数字还必须制定，这个值就是刚好可以满足全部宽度
                doc.add_paragraph("\n")
                dl2 = doc.add_paragraph(nr2)
                pf2 = dl2.paragraph_format
                pf2.first_line_indent = Inches(0.3)
                # 设置段落格式
                pf2.space_before = Pt(5)  # 上行间距
                pf2.space_after = Pt(10)  # 下行间距
                pf2.line_spacing = Pt(20)  # 行距
                # 处理下级循环内容
                for ej in ml.get(dl):
                    # for ejbt in ej.keys():
                    #     # print("第二级标题",ejbt)
                    p2 = doc.add_heading(level=2)
                    # p2.style=style
                    p2ejbt = p2.add_run(ej.keys())  # 直接输出
                    p2ejbt.font.name = u'宋体'
                    p2ejbt._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
                    p2ejzt = p2ejbt.font
                    p2ejzt.size = Pt(15)
                    # 设置段落格式
                    p2ejbt.space_before = Pt(5)  # 上行间距
                    p2ejbt.space_after = Pt(10)  # 下行间距
                    p2ejbt.line_spacing = Pt(20)  # 行距
                    # p2ejzt.bold=True
                    p2ejzt.color.rgb = RGBColor(10, 10, 10)
                    for ejz in ej.values():
                        for sj in ejz:
                            p23 = doc.add_heading(level=3)
                            # p23.style=style
                            p23bt = p23.add_run(sj.keys())  # 直接输出
                            p23bt.font.name = u'宋体'
                            p23bt._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
                            p23zt = p23bt.font
                            p23zt.size = Pt(13)
                            # p23zt.bold = True
                            p23zt.color.rgb = RGBColor(10, 10, 10)
                            # 直接输出内容
                            paragraph = doc.add_paragraph(sj.values())
                            paragraph_format = paragraph.paragraph_format
                            paragraph_format.first_line_indent = Inches(0.3)
                            paragraph_format.space_after = Pt(10)  # 下行间距
                            paragraph_format.space_before = Pt(5)  # 上行间距
                            paragraph_format.line_spacing = Pt(20)  # 行距

                # 第二章最后增加总结的一段话
                paragraph = doc.add_paragraph(zj_2)
                paragraph_format = paragraph.paragraph_format
                paragraph_format.first_line_indent = Inches(0.3)
                paragraph_format.space_before = Pt(5)  # 上行间距
                paragraph_format.space_after = Pt(10)  # 下行间距
                paragraph_format.line_spacing = Pt(20)  # 行距


            else:
                # 处理除了第二章之外的规律结构
                for xl in ml.get(dl):
                    # 每个小类中的键值
                    for ej in xl.keys():
                        # xldl = doc.add_paragraph()  # 小类段落
                        # xlbt = xldl.add_run(ej[:ej.rindex("\t")])
                        # xlbtzt = xlbt.font
                        # xlbtzt.size = Pt(13)
                        # xlbtzt.bold = True
                        xldl = doc.add_heading(level=2)  # 添加标题
                        # xldl.style=style
                        xlbt = xldl.add_run(ej[:ej.rindex("\t")])
                        xlbt.font.name = u'宋体'
                        xlbt._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
                        xlbtzt = xlbt.font
                        # xlbtzt.bold=True
                        xlbtzt.size = Pt(15)
                        xlbtzt.color.rgb = RGBColor(0, 0, 0)

                        # 获取到每一个二级下面的内容
                        num = 0  # 定义一个数量用来判断循环内获取到第几个内容，来方便插入图表
                        for nr in xl.get(ej):
                            num += 1
                            paragraph = doc.add_paragraph(nr)
                            paragraph_format = paragraph.paragraph_format
                            paragraph_format.first_line_indent = Inches(0.3)
                            paragraph_format.space_before = Pt(5)  # 上行间距
                            paragraph_format.space_after = Pt(10)  # 下行间距
                            paragraph_format.line_spacing = Pt(20)  # 行距

                            # ############### 处理有图表的特殊部分 ################
                            sql_dx = getsqldata()

                            if ej == "3.2\t分发库运行情况\t\t" and num == 1:
                                # print(sql.get_b32_ycjc())
                                data = getsqldata.getdata(sql_dx, sql.get_b32_ycjc())
                                table = doc.add_table(rows=1, cols=3, style=table_style)
                                hdr_cells = table.rows[0].cells
                                hdr_cells[0].text = "日期"
                                hdr_cells[1].text = "异常进程数"
                                hdr_cells[2].text = "异常进程单位"
                                hdr_cells[0].width = Inches(1)
                                hdr_cells[1].width = Inches(0.3)
                                hdr_cells[2].width = Inches(4.5)
                                # 设置标题居中
                                for i in range(3):
                                    cp = table.cell(0, i).paragraphs[0]
                                    cp.paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

                                day = []
                                ycs = []
                                dw = set()  # 定义单位为集合方便去重
                                for rq, ycjcs, ycjcdw in data:
                                    row_cells = table.add_row().cells
                                    row_cells[0].text = rq
                                    row_cells[1].text = str(ycjcs)  # 不转为字符串不认识
                                    row_cells[2].text = ycjcdw
                                    # print("{rq}>>>{jcs}>>>{ycdw}".format(rq=rq, jcs=ycjcs, ycdw=ycjcdw))
                                    day.append(int(rq[-2:]))
                                    ycs.append(ycjcs)
                                    if ycjcdw:  # 过滤掉为空的
                                        dw.add(ycjcdw)
                                doc.add_paragraph("\n")
                                # 添加折线图
                                bt = "图3.2.1 {ny}分日分发库异常进程情况统计图".format(ny=self.ny)
                                # l321 = Line()
                                # l321.add_xaxis(day).add_yaxis("异常数", ycs).set_global_opts(title_opts=opts.TitleOpts(title=bt))
                                # l321 = Line(init_opts=opts.InitOpts(width="800px", height="400px"))
                                # (
                                #
                                #     l321.add_xaxis(xaxis_data=day)
                                #         .add_yaxis(
                                #         series_name="最高气温",
                                #         y_axis=ycs,
                                #         markpoint_opts=opts.MarkPointOpts(
                                #             data=[
                                #                 opts.MarkPointItem(type_="max", name="最大值"),
                                #                 opts.MarkPointItem(type_="min", name="最小值"),
                                #             ]
                                #         ),
                                #         markline_opts=opts.MarkLineOpts(
                                #             data=[opts.MarkLineItem(type_="average", name="平均值")]
                                #         ),
                                #     )
                                #         .set_global_opts(
                                #         title_opts=opts.TitleOpts(title=bt, subtitle=self.ny)
                                #     )
                                # )
                                bar321 = Bar()
                                (

                                    bar321.add_xaxis(day)
                                        .add_yaxis("异常数", ycs, color="#019CF6")
                                        # .reversal_axis()
                                        .set_series_opts(label_opts=opts.LabelOpts(position="right"))
                                        .set_global_opts(title_opts=opts.TitleOpts(title="分发库异常展示-分日期"))
                                        .render("temp.html")
                                )

                                make_snapshot(snapshot, bar321.render(), r"..\img\img_nr3201_ffkyc_frq.png")
                                self.doc.add_picture(r"..\img\img_nr3201_ffkyc_frq.png", width=Inches(6.25),
                                                     height=Inches(4))

                                # 添加总结概括的一段
                                data = getsqldata.getdata(sql_dx, sql.get_zj32_ffkyc_frq())
                                # print(data)
                                str_yc = ""
                                for mx in data:
                                    str_yc += reduce(lambda a, b: "{a}有{b}次异常，".format(a=a, b=b), mx)
                                temp_nr = "分日期统计可以看出{yf}月份从1号到{end}号每天的任务运行失败情况，其中失败任务依次是:{str_yc}{yf}月共有{wycts}天没有失败的任务。".format(
                                    yf=yf, end=end, str_yc=str_yc, wycts=end - len(data))
                                p = self.doc.add_paragraph(temp_nr)
                                pf = p.paragraph_format
                                pf.first_line_indent = Inches(0.3)
                                pf.space_before = Pt(5)  # 上行间距
                                pf.space_after = Pt(10)  # 下行间距
                                pf.line_spacing = Pt(20)  # 行距包含当前行字体在内的距离

                                # 图2 任务分单位异常情况
                                data = getsqldata.getdata(sql_dx, sql.get_b32_ffkyc_fdw())
                                dd = defaultdict(lambda: 0)

                                # 处理地图需要的数据以及增加表格其他列
                                for mx in data:
                                    dd[mx[0]] = mx[1]

                                self.doc.add_paragraph("\n按地图显示如下：")

                                # 地图部分
                                m1 = Map()
                                m1.add(series_name="分发库异常", data_pair=data, maptype="china",
                                       name_map={"key": "value"},
                                       is_map_symbol_show=True)
                                m1.set_global_opts(
                                    title_opts=opts.TitleOpts(title="图3.2.2 分发库异常情况分单位", subtitle=self.ny),
                                    # 设置标题和副标题在左侧
                                    visualmap_opts=opts.VisualMapOpts(is_piecewise=True, max_=max(dd.values()))
                                    # 是否分段,以及分段的最大值
                                )

                                make_snapshot(snapshot, m1.render(), r'..\img\img_nr32_ffkycdw.png')
                                self.doc.add_picture(r'..\img\img_nr32_ffkycdw.png', width=Inches(6.25))

                                # 添加总结概括的一段
                                yczs = sum(ycs)
                                dws = len(dw)
                                # reduce(function, sequence[, initial] ) -> value
                                # function参数是一个有两个参数的函数，reduce依次从sequence中取一个元素，和上一次调用function的结果做参数再次调用function。
                                # 第一次调用function时，如果提供initial参数，会以sequence中的第一个元素和initial作为参数调用function，否则会以序列sequence中的前两个元素做参数调用function。
                                str_yc = ""
                                for mx in data:
                                    str_yc += reduce(lambda a, b: "{a}{b}次异常".format(a=a, b=b), mx)
                                temp_nr = "如上图所示，{ny}整月分发库OGG链路进程共出现{yczs}次异常，共涉及{dws}家单位，分别是：{str_yc}。".format(
                                    ny=self.ny, yczs=yczs, dws=dws, str_yc=str_yc)
                                p = self.doc.add_paragraph(temp_nr)
                                pf = p.paragraph_format
                                pf.first_line_indent = Inches(0.3)
                                pf.space_before = Pt(5)  # 上行间距
                                pf.space_after = Pt(10)  # 下行间距
                                pf.line_spacing = Pt(20)  # 行距

                            # 3.3 ############################################
                            if ej == "3.3\t分发库异常分析\t\t" and num == 1:
                                data = getsqldata.getdata(sql_dx, sql.get_zj32_ffkyc_frq())
                                # print(data)
                                table = doc.add_table(rows=3, cols=3, style=table_style)
                                table.alignment = WD_TABLE_ALIGNMENT.CENTER  # 设置表格居中
                                # 注意：表的所有列宽度合计为10，所以在设置表格每列宽宽度时要同时设置所有列宽，并且合计为10。如果只设置某一列宽，那么其余列将平分剩余宽度。如果只设置某几列宽，将不起作用。
                                table.cell(0, 1).width = Inches(5.4)
                                hdr_cells = table.rows[0].cells
                                hdr_cells[0].text = "序号"
                                hdr_cells[1].text = "OGG进程异常原因"
                                hdr_cells[2].text = "次数"

                                # 设置标题居中
                                for i in range(3):
                                    cp = table.cell(0, i).paragraphs[0]
                                    cp.paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

                                ycyy = ['人为导致：未按版本发布规范进行操作导致进程异常', '系统故障：分发库机器检修及扩容']  # 名称对象
                                yccs = []  # 数据对象
                                explode = []  # 爆裂值
                                for yy, cs in data:
                                    yccs.append(cs)
                                    explode.append(0)
                                yczs = str(sum(yccs))
                                # 此表格暂时约定为这样，所以此处除了错误数量，其他是固定的
                                table.cell(1, 0).text = "1"
                                table.cell(1, 1).text = ycyy[0]
                                table.cell(1, 2).text = yczs
                                table.cell(2, 0).text = "2"
                                table.cell(2, 1).text = ycyy[1]
                                table.cell(2, 2).text = "0"

                                temp_nr = "如上图所示：整个{yf}月份分发库OGG进程异常共有{zcs}次，其中人为原因导致异常的有{zcs}次，通过对异常分析,得出产生问题的主要原因是人为导致的，问题依旧集中在分发库版本升级后未按规定步骤刷新定义文件导致解析进程异常。".format(
                                    yf=yf, zcs=sum(yccs))
                                doc.add_paragraph("\n")
                                p = self.doc.add_paragraph(temp_nr)
                                pf = p.paragraph_format
                                pf.first_line_indent = Inches(0.3)
                                pf.space_before = Pt(5)  # 上行间距
                                pf.space_after = Pt(10)  # 下行间距
                                pf.line_spacing = Pt(20)  # 行距
                                # 首行缩进0.74厘米，即2个字符上面那个英寸还真不好掌控
                                # paragraph_format.first_line_indent = Cm(0.74)

                            # 4.1 ############################################
                            if ej == "4.1\t云平台数据集成概述\t\t" and num == 1:
                                data = getsqldata.getdata(sql_dx, sql.get_b41_yptjcbsqk())
                                # print(data)
                                # 数据更新情况表
                                table = doc.add_table(rows=2, cols=8, style=table_style)
                                table.cell(0, 0).text = "分类"
                                table.cell(0, 0).merge(table.cell(1, 0))
                                table.cell(0, 1).text = "来源系统"
                                table.cell(0, 1).merge(table.cell(1, 1))
                                table.cell(0, 1).width = Inches(2)
                                table.cell(0, 2).text = "集成/推送表范围"
                                table.cell(0, 2).merge(table.cell(0, 3))
                                table.cell(1, 2).text = "镜像层"
                                table.cell(1, 3).text = "基础层"
                                table.cell(0, 4).text = "任务部署情况"
                                table.cell(0, 4).merge(table.cell(0, 5))
                                table.cell(1, 4).text = "镜像层"
                                table.cell(1, 5).text = "基础层"
                                table.cell(0, 6).text = "同步频率"
                                table.cell(0, 6).merge(table.cell(1, 6))
                                table.cell(0, 7).text = "集成/推送方式"
                                table.cell(0, 7).merge(table.cell(1, 7))
                                # 处理每一行
                                row_num = 1  # 略过前两行表头
                                # 总结段落1数据
                                temp_list1 = [data[0][1]]  # 预加载第一个值
                                dict1 = {data[0][0]: temp_list1}  # 初始字典
                                for mx in data:
                                    row_cells = table.add_row().cells
                                    row_num += 1
                                    for i in range(len(mx)):
                                        row_cells[i].text = str(mx[i])
                                    # 处理首列合并的问题
                                    # print(row_num)
                                    if row_num > 2:
                                        # 如果和前面内容一致则进行合并单元格，临时列表一直拼接值，否则的话字典追加新对象,并且临时列表清空追加当前值，重新开始插入新数据
                                        if data[row_num - 2][0] == data[row_num - 3][0]:
                                            table.cell(row_num - 1, 0).text = ''
                                            table.cell(row_num - 1, 0).merge(table.cell(row_num, 0))
                                            temp_list1.append(data[row_num - 2][1])  # 拼接当前值
                                        else:
                                            temp_list1 = [data[row_num - 2][1]]  # 临时列表清空追加当前值
                                            dict1[data[row_num - 2][0]] = temp_list1  # 更换新的主键

                                # 设置表所有内容居中
                                for row_num in range(len(data) + 2):  # 要加上标题的两行
                                    for col_num in range(8):
                                        cp = table.cell(row_num, col_num).paragraphs[0]
                                        cpf = cp.paragraph_format
                                        cpf.alignment = WD_TABLE_ALIGNMENT.CENTER
                                        cpf.alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

                                # 动态生成底下说明段落内容,第1，2动态获取，其他固定值
                                temp_nr = "1、云平台现已集成的数据有："
                                for a in dict1.keys():
                                    # 如果是多个需要取消合计
                                    if a != "数据推送同步":
                                        if len(dict1.get(a)) > 1:
                                            temp_list1 = dict1.get(a)[:-1]
                                        else:
                                            temp_list1 = dict1.get(a)
                                    temp_nr = temp_nr + a + "包含" + ",".join(temp_list1) + "；"

                                self.doc.add_paragraph("\n")

                                p = self.doc.add_paragraph(temp_nr)
                                pf = p.paragraph_format
                                pf.first_line_indent = Inches(0.3)
                                pf.space_before = Pt(5)  # 上行间距
                                pf.space_after = Pt(10)  # 下行间距
                                pf.line_spacing = Pt(20)  # 行距

                                temp_list1 = dict1.get("数据推送同步")[:-1]
                                temp_nr = "2、云平台向外部推送及同步的数据有：" + ",".join(temp_list1)
                                p = self.doc.add_paragraph(temp_nr)
                                pf = p.paragraph_format
                                pf.first_line_indent = Inches(0.3)
                                pf.space_before = Pt(5)  # 上行间距
                                pf.space_after = Pt(10)  # 下行间距
                                pf.line_spacing = Pt(20)  # 行距

                            # 4.2 ############################################
                            if ej == "4.2\t数据集成任务情况\t\t" and num == 1:
                                # 增加总结段落
                                data = getsqldata.getdata(sql_dx, sql.get_zj421_xzrwqk())
                                # print(data)
                                rwzs, dcrws, bsl = data[0]
                                temp_nr = "从{yf}月1号到{yf}月{end}号总共新增了{rwzs}个任务（基础层新增{bsl}张表，部署任务{dcrws}个；镜像层新增{bsl}张表，部署任务{dcrws}个任务）。".format(
                                    yf=yf, end=end, rwzs=rwzs, dcrws=dcrws, bsl=bsl)
                                p = self.doc.add_paragraph(temp_nr)
                                pf = p.paragraph_format
                                pf.first_line_indent = Inches(0.3)
                                pf.space_before = Pt(5)  # 上行间距
                                pf.space_after = Pt(10)  # 下行间距
                                pf.line_spacing = Pt(20)  # 行距

                                # 按时间总结的一段话
                                data = getsqldata.getdata(sql_dx, sql.get_zj422_rwyc_frq())

                                # p4301 = Pie()
                                # c = (
                                #
                                #     p4301.add(
                                #         "",
                                #         [list(z) for z in data],
                                #         center=["50%", "50%"],  # 圆心的横竖坐标位置
                                #     )
                                #         .set_series_opts(label_opts=opts.LabelOpts(formatter="{b}: {c}"))
                                # )
                                # 数据太多饼图不好看，还是使用折线图
                                bt = "图4.2.1 {ny}任务异常情况按日统计图".format(ny='202006')
                                day = []
                                ycs = []
                                for mx in data:
                                    day.append(mx[0][-2:])
                                    ycs.append(mx[1])

                                zxt = Line()
                                zxt.add_xaxis(day).add_yaxis("异常数", ycs).set_global_opts(
                                    title_opts=opts.TitleOpts(title=bt))
                                zxt = Line(init_opts=opts.InitOpts(width="800px", height="400px"))
                                (

                                    zxt.add_xaxis(xaxis_data=day)
                                        .add_yaxis(
                                        series_name="",
                                        y_axis=ycs,
                                        # symbol="triangle",
                                        # symbol_size=20,
                                        linestyle_opts=opts.LineStyleOpts(color="green", width=4, type_="dashed"),
                                        # label_opts=opts.LabelOpts(is_show=False),
                                        # itemstyle_opts=opts.ItemStyleOpts(
                                        #     border_width=3, border_color="yellow", color="blue"
                                        # ),
                                        markpoint_opts=opts.MarkPointOpts(
                                            data=[
                                                opts.MarkPointItem(type_="max", name="最大值"),
                                                opts.MarkPointItem(type_="min", name="最小值"),
                                            ]
                                        ),
                                        markline_opts=opts.MarkLineOpts(
                                            data=[opts.MarkLineItem(type_="average", name="平均值")]
                                        ),
                                    )
                                        .set_global_opts(
                                        title_opts=opts.TitleOpts(title=bt, subtitle=self.ny)
                                    )
                                )

                                # zxt.render("temp.html")

                                make_snapshot(snapshot, zxt.render(), r"..\img\img_nr4202_rwyc_frq.png")
                                self.doc.add_picture(r"..\img\img_nr4202_rwyc_frq.png", width=Inches(6.25))

                                data = getsqldata.getdata(sql_dx, sql.get_zj421_rwyyc_frq())
                                str_yc = ""
                                for mx in data[:3]:
                                    str_yc += reduce(lambda a, b: "{yf}月{a}号有{b}次异常,".format(yf=yf, a=int(a[-2:]), b=b),
                                                     mx)
                                temp_nr = "分日期统计可以看出{yf}月份从1号到{end}号每天的任务运行失败情况，其中失败任务数排在前三的依次是:{str_yc}。{yf}月共有{wycts}天没有失败的任务。".format(
                                    yf=yf, end=end, str_yc=str_yc, wycts=end - len(data))
                                p = self.doc.add_paragraph(temp_nr)
                                pf = p.paragraph_format
                                pf.first_line_indent = Inches(0.3)
                                pf.space_before = Pt(5)  # 上行间距
                                pf.space_after = Pt(10)  # 下行间距
                                pf.line_spacing = Pt(20)  # 行距

                                p = doc.add_paragraph("{nf}年{yf}月任务分单位异常情况：".format(nf=nf, yf=yf))
                                pf = p.paragraph_format
                                pf.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                                # 任务分单位异常情况
                                data = getsqldata.getdata(sql_dx, sql.get_b42_rwyc_fdw())
                                dd = defaultdict(lambda: 0)

                                temp_nr = ""
                                # 生成表格，因为地图也表示不全
                                table = doc.add_table(rows=1, cols=4, style=table_style)
                                table.width = Inches(4)
                                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                                hdr_cells = table.rows[0].cells
                                hdr_cells[0].text = "序号"
                                hdr_cells[1].text = "所属月份"
                                hdr_cells[2].text = "单位名称"
                                hdr_cells[3].text = "异常次数"
                                # 处理地图需要的数据以及增加表格其他列
                                temp_num = 0
                                for mx in data:
                                    temp_num += 1
                                    dd[mx[0]] = mx[1]
                                    row_cells = table.add_row().cells
                                    row_cells[0].text = str(temp_num)
                                    row_cells[1].text = self.ny
                                    row_cells[2].text = mx[0]
                                    row_cells[3].text = str(mx[1])

                                self.doc.add_paragraph("\n按地图显示如下：")

                                # 地图部分
                                m1 = Map()
                                m1.add(series_name="异常情况", data_pair=data, maptype="china",
                                       name_map={"key": "value"},
                                       is_map_symbol_show=True)
                                m1.set_global_opts(
                                    title_opts=opts.TitleOpts(title="图4.2.2 异常情况分单位", subtitle=self.ny),
                                    # 设置标题和副标题在左侧
                                    visualmap_opts=opts.VisualMapOpts(is_piecewise=True, max_=max(dd.values()))
                                    # 是否分段,以及分段的最大值
                                )

                                # m1.render("temp.html")
                                # make_snapshot(snapshot,file_name=r'temp.html',output_name=r'..\img\img_nr32_ycfdw.png')

                                # from snapshot_selenium import snapshot
                                # 想要使用他必须安装驱动，谷歌浏览器驱动的下载网址http://chromedriver.storage.googleapis.com/index.html
                                # 只需要把它放到下面两个路径下即可!
                                # 还有一个就是在你python安装的目录下,
                                #  一个谷歌浏览器的位置,具体的位置:C:\Program Files (x86)\Google\Chrome\Application
                                make_snapshot(snapshot, m1.render(), r'..\img\img_nr42_rwycfdw.png')

                                self.doc.add_picture(r'..\img\img_nr42_rwycfdw.png', width=Inches(6.25))

                                yczs = sum(dd.values())
                                dws = len(dd.keys())
                                str_dw = reduce(lambda a, b: "{a},{b}".format(a=a, b=b), dd.keys())
                                temp_nr = "如上图所示，{ny}整月集成问题共出现{yczs}次，共涉及{dws}家单位，分别是：{str_dw}。".format(
                                    ny=self.ny, yczs=yczs, dws=dws, str_dw=str_dw)
                                p = self.doc.add_paragraph(temp_nr)
                                pf = p.paragraph_format
                                pf.first_line_indent = Inches(0.3)
                                pf.space_before = Pt(5)  # 上行间距
                                pf.space_after = Pt(10)  # 下行间距
                                pf.line_spacing = Pt(20)  # 行距

                            # 4.3 ############################################
                            if ej == "4.3\t运行失败任务分析\t\t" and num == 1:

                                # 饼图 任务异常原因分析
                                data = getsqldata.getdata(sql_dx, sql.get_b43_rwyc_fyy())
                                yccs = []
                                ycyy = []
                                explode = []
                                for mx in data:
                                    yccs.append(mx[1])
                                    ycyy.append(mx[0])
                                    explode.append(0)
                                # 如果需要的情况增加上其他
                                if int(yczs) - sum(yccs) > 0:
                                    yccs.append()
                                    ycyy.append("其他原因")
                                    explode.append(0)

                                temp_nr = "图4.3.1  任务主要异常原因分析"
                                doc.add_paragraph("\n")
                                p = doc.add_paragraph()
                                temp_title = p.add_run(temp_nr)
                                font = temp_title.font
                                font.bold = True
                                pf = p.paragraph_format
                                pf.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                # 如果有数据增加异常原因饼图
                                if len(yccs) > 0:
                                    explode[0] = 0.1  # 改变自己需要的分离缝隙  各个值所属模块突出来的缝隙大小

                                    # matplotlib 的饼状图
                                    # fig1, ax1 = plt.subplots()
                                    # ax1.pie(yccs, explode=explode, labels=ycyy, autopct='%1.2f%%',
                                    #         shadow=True, startangle=90)  # 分别是数据对象，爆裂值，名称列表，自动百分比，是否阴影，起始角度
                                    # ax1.axis('equal')  # Equal aspect ratio ensures that pie is drawn as a circle.
                                    # plt.savefig(r"..\img\img_nr4301_rwyc_fyy.png")
                                    # # plt.show()
                                    p4301 = Pie()
                                    c = (

                                        p4301.add(
                                            "",
                                            [list(z) for z in zip(ycyy, yccs)],
                                            center=["50%", "50%"],  # 圆心的横竖坐标位置
                                        )
                                            .set_global_opts(
                                            title_opts=opts.TitleOpts(title="任务异常-分原因"),
                                            legend_opts=opts.LegendOpts(pos_left="15%"),  # 上面每一项距离指定一侧的距离
                                        )
                                            .set_series_opts(label_opts=opts.LabelOpts(formatter="{b}: {c}"))
                                        # .render("pie_position.html")
                                    )
                                    make_snapshot(snapshot, p4301.render(), r"..\img\img_nr4301_rwyc_fyy.png")
                                    self.doc.add_picture(r"..\img\img_nr4301_rwyc_fyy.png", width=Inches(6.25))

                                # 添加总结概括的一段
                                if len(ycyy) > 0:
                                    str_yc = "错误率比较高的是：" + reduce(lambda a, b: "{a},{b}".format(a=a, b=b), ycyy)
                                else:
                                    str_yc = ""

                                temp_nr = "如上图所示，{yf}月集成问题共出现{yczs}次，{str_yc}。导致失败的原因主要有三大类，一是分发库问题，主要是连接异常，版本不一致，脏数据（最大比重的就是脏数据问题）；二个是云平台问题，主要是实例调度问题、任务卡住、资源不足导致任务失败的问题。从统计数量来看脏数据问题导致的失败任务是最多的，这个问题需要各地运维人员严格按照规范来操作，也是后期需要解决的问题。".format(
                                    yf=yf, yczs=yczs, str_yc=str_yc)
                                p = self.doc.add_paragraph(temp_nr)
                                pf = p.paragraph_format
                                pf.first_line_indent = Inches(0.3)
                                pf.space_before = Pt(5)  # 上行间距
                                pf.space_after = Pt(10)  # 下行间距
                                pf.line_spacing = Pt(20)  # 行距

                            # 4.4 ############################################
                            if ej == "4.4\t失败任务应对策略\t\t" and num == 1:
                                data = getsqldata.getdata(sql_dx, sql.get_b44_sbrwydcl())
                                # print(data)
                                table = doc.add_table(rows=1, cols=3, style=table_style)
                                hdr_cells = table.rows[0].cells
                                hdr_cells[0].text = "序号"
                                hdr_cells[0].width = Cm(2)
                                hdr_cells[1].text = "失败类别"
                                hdr_cells[2].text = "解决方法"
                                for mx in data:
                                    row_cells = table.add_row().cells
                                    row_cells[0].text = str(int(mx[0]))
                                    row_cells[1].text = mx[1]
                                    row_cells[2].text = mx[2]

                                # table.autofit=True
                                # hdr_cells[0].width = Inches(0.5)

                            # 4.5 生成数据一致性Excel
                            if ej == "4.5\t数据集成一致性监控处理情况\t\t":
                                if num == 1:
                                    # 添加前提说明
                                    temp_nr = "以下选取5张金三核心表对{yf}月24号至{yf}月28号5天数据抽取情况进行数据量分析比对，详情见附件：《{ny}数据抽样比对结果.xlsx》".format(
                                        yf=yf, ny=self.ny)
                                    p = self.doc.add_paragraph(temp_nr)
                                    pf = p.paragraph_format
                                    pf.first_line_indent = Inches(0.3)
                                    pf.space_before = Pt(5)  # 上行间距
                                    pf.space_after = Pt(10)  # 下行间距
                                    pf.line_spacing = Pt(20)  # 行距

                                    list_tablename = ['HX_DJ_DJ_NSRXX', 'HX_DJ_DJ_NSRXX_KZ', 'HX_RD_RD_SFZRDXXB',
                                                      'HX_SB_SB_SBXX', 'HX_ZS_ZS_YJSF']
                                    workbook = xlsxwriter.Workbook(r"..\appendix\{ny}数据抽样比对结果.xlsx".format(ny=self.ny))
                                    for tablename in list_tablename:
                                        worksheet = workbook.add_worksheet(tablename)
                                        # 声明一个加粗的样式用来增加给表头，颜色是16进制数
                                        bt = workbook.add_format(
                                            {'bold': True, 'bg_color': '9BC2E6', 'align': 'centre', 'valign': 'vcenter',
                                             'border': 1})
                                        nr = workbook.add_format({'align': 'centre', 'valign': 'vcenter', 'border': 1})
                                        worksheet.merge_range(first_row=0, last_row=1, first_col=0, last_col=0,
                                                              data="单位", cell_format=bt)
                                        worksheet.merge_range(first_row=0, last_row=0, first_col=1, last_col=4,
                                                              data="{nf}年{yf}月24日".format(nf=nf, yf=yf), cell_format=bt)
                                        worksheet.merge_range(first_row=0, last_row=0, first_col=5, last_col=8,
                                                              data="{nf}年{yf}月25日".format(nf=nf, yf=yf), cell_format=bt)
                                        worksheet.merge_range(first_row=0, last_row=0, first_col=9, last_col=12,
                                                              data="{nf}年{yf}月26日".format(nf=nf, yf=yf), cell_format=bt)
                                        worksheet.merge_range(first_row=0, last_row=0, first_col=13, last_col=16,
                                                              data="{nf}年{yf}月27日".format(nf=nf, yf=yf), cell_format=bt)
                                        worksheet.merge_range(first_row=0, last_row=0, first_col=17, last_col=20,
                                                              data="{nf}年{yf}月28日".format(nf=nf, yf=yf), cell_format=bt)
                                        worksheet.write(1, 1, "分发库", bt)
                                        worksheet.write(1, 2, "基础层", bt)
                                        worksheet.write(1, 3, "差异", bt)
                                        worksheet.write(1, 4, "差异率", bt)
                                        worksheet.write(1, 5, "分发库", bt)
                                        worksheet.write(1, 6, "基础层", bt)
                                        worksheet.write(1, 7, "差异", bt)
                                        worksheet.write(1, 8, "差异率", bt)
                                        worksheet.write(1, 9, "分发库", bt)
                                        worksheet.write(1, 10, "基础层", bt)
                                        worksheet.write(1, 11, "差异", bt)
                                        worksheet.write(1, 12, "差异率", bt)
                                        worksheet.write(1, 13, "分发库", bt)
                                        worksheet.write(1, 14, "基础层", bt)
                                        worksheet.write(1, 15, "差异", bt)
                                        worksheet.write(1, 16, "差异率", bt)
                                        worksheet.write(1, 17, "分发库", bt)
                                        worksheet.write(1, 18, "基础层", bt)
                                        worksheet.write(1, 19, "差异", bt)
                                        worksheet.write(1, 20, "差异率", bt)
                                        # 获取数据进行添加
                                        # print(sql.get_b45_sjyzx(tablename=tablename))
                                        data = getsqldata.getdata(sql_dx, sql.get_b45_sjyzx(tablename=tablename))
                                        for row_num in range(len(data)):
                                            for col_num in range(len(data[0])):
                                                worksheet.write(row_num + 2, col_num, data[row_num][col_num], nr)
                                    workbook.close()
                                    # 将生成好的Excel附加进来
                                elif num == 3:
                                    # 生成 4.5 excel 2 数据差异运维记录
                                    workbook = xlsxwriter.Workbook(r"..\appendix\{ny}数据差异运维记录.xlsx".format(ny=self.ny))
                                    worksheet = workbook.add_worksheet("运维记录")
                                    # 声明一个加粗的样式用来增加给表头，颜色是16进制数
                                    bt = workbook.add_format({
                                        'bold': True, 'bg_color': '0197F5', 'align': 'centre', 'valign': 'vcenter',
                                        'border': 1
                                    })
                                    nr = workbook.add_format({'align': 'centre', 'valign': 'vcenter',
                                                              'border': 1})
                                    worksheet.write(0, 0, "项目名称", bt)
                                    worksheet.write(0, 1, "业务日期", bt)
                                    worksheet.write(0, 2, "异常任务名称", bt)
                                    worksheet.write(0, 3, "问题类型", bt)
                                    worksheet.write(0, 4, "解决方案", bt)
                                    # 设置列宽
                                    worksheet.set_column(0, 1, 20)
                                    worksheet.set_column(2, 4, 40)

                                    data = getsqldata.getdata(sql_dx, sql.get_b45_sjcyyw())
                                    for row_num in range(len(data)):
                                        for col_num in range(len(data[0])):
                                            worksheet.write(row_num + 1, col_num, data[row_num][col_num], nr)
                                    workbook.close()
                                    # 插入word

                            # 4.6 ############################################
                            if ej == "4.6\t源端数据更新情况\t\t" and num == 1:
                                data = getsqldata.getdata(sql_dx, sql.get_b46_sjgxqk())
                                # print(data)
                                # 数据更新情况表
                                table = doc.add_table(rows=1, cols=5, style=table_style)
                                hdr_cells = table.rows[0].cells
                                hdr_cells[0].text = "月份"
                                hdr_cells[1].text = "数据来源"
                                hdr_cells[1].width = Cm(6)
                                hdr_cells[2].text = '总量'
                                hdr_cells[3].text = "有更新表数量"
                                hdr_cells[4].text = "未更新表数量"
                                # 设置标题居中
                                for i in range(5):
                                    cp = table.cell(0, i).paragraphs[0]
                                    cp.paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                                # 定义5个列表给统计用，为了提高垃圾回收机制，我们临时的都使用统一规范
                                temp_list4 = []
                                temp_list1 = []
                                temp_list2 = []
                                temp_list3 = []
                                for mx in data:
                                    row_cells = table.add_row().cells
                                    row_cells[0].text = mx[0]
                                    row_cells[1].text = mx[1]
                                    row_cells[2].text = str(mx[2])
                                    row_cells[3].text = str(mx[3])
                                    row_cells[4].text = str(mx[4])

                                    temp_list4.append(mx[4])
                                    temp_list1.append(mx[1])
                                    temp_list2.append(mx[2])
                                    temp_list3.append(mx[3])

                                doc.add_paragraph("\n")

                                temp_nr = "目前（{yf}）云平台基础层配置调度任务的数据源共有{lys}个，包含日调度、周调度、月调度三种更新频率，共计{zbs}张表。根据本月监控源端数据变化情况统计，其中按调度更新的表有{gxs}张，未更新的表有{wgxs}张。各源端系统详细情况如上表所示:".format(
                                    yf=self.ny, lys=len(temp_list1), zbs=sum(temp_list2), gxs=sum(temp_list3),
                                    wgxs=sum(temp_list4))
                                p = doc.add_paragraph(temp_nr)
                                pf = p.paragraph_format
                                pf.first_line_indent = Inches(0.3)
                                pf.space_before = Pt(5)  # 上行间距
                                pf.space_after = Pt(10)  # 下行间距
                                pf.line_spacing = Pt(20)  # 行距

            doc.add_page_break()

        doc.save(self.wdmc)


if __name__ == '__main__':
    start_time = time.time()
    scword("202006").scword()
    # scword("202006").checkdir()
    print("花费秒数：", time.time() - start_time)
